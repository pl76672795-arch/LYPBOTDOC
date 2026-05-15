import logging
import asyncio
from pathlib import Path
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

logger = logging.getLogger(__name__)

class GeneradorWord:
    """Motor profesional para la generación de documentos Word mediante plantillas."""
    
    def __init__(self, template_path: str | Path):
        self.template_path = Path(template_path)
        
    def _crear_plantilla_automatica(self):
        """Si la plantilla no existe, crea una perfecta automáticamente según el nombre."""
        if self.template_path.exists():
            return
            
        self.template_path.parent.mkdir(parents=True, exist_ok=True)
        
        nombre_archivo = self.template_path.name.lower()
        if "demanda" in nombre_archivo:
            self._crear_plantilla_demanda()
        elif "compraventa" in nombre_archivo:
            self._crear_plantilla_compraventa()
        elif "liquidacion" in nombre_archivo:
            self._crear_plantilla_liquidacion()
        elif "certificado" in nombre_archivo:
            self._crear_plantilla_certificado()
        else:
            self._crear_plantilla_arrendamiento()

    def _crear_plantilla_arrendamiento(self):
        doc = Document()
        
        # --- ESTÉTICA PREMIUM ---
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Cm(2.54)
            section.left_margin = section.right_margin = Cm(2.54)
            
        def _add_justified(text=''):
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return p

        doc.add_heading('CONTRATO DE ARRENDAMIENTO DE BIEN INMUEBLE', 1)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = _add_justified()
        p.add_run('Conste por el presente documento privado, el Contrato de Arrendamiento de Bien Inmueble que celebran, de una parte, don(ña) ')
        p.add_run('{{ nombre_arrendador }}').bold = True
        p.add_run(', identificado(a) con DNI N° ')
        p.add_run('{{ dni_arrendador }}').bold = True
        p.add_run(', a quien en adelante se le denominará EL ARRENDADOR; y de la otra parte, don(ña) ')
        p.add_run('{{ nombre_arrendatario }}').bold = True
        p.add_run(', identificado(a) con DNI N° ')
        p.add_run('{{ dni_arrendatario }}').bold = True
        p.add_run(', a quien en adelante se le denominará EL ARRENDATARIO; en los términos y condiciones contenidos en las siguientes cláusulas:\n')

        doc.add_heading('PRIMERA: ANTECEDENTES', 2)
        p1 = _add_justified()
        p1.add_run('EL ARRENDADOR es legítimo propietario del inmueble ubicado en ')
        p1.add_run('{{ direccion_inmueble }}').bold = True
        p1.add_run(', el mismo que se encuentra en perfecto estado de conservación, con sus instalaciones eléctricas, sanitarias y de agua operativas y al día en sus pagos.')

        doc.add_heading('SEGUNDA: OBJETO DEL CONTRATO', 2)
        _add_justified('Por el presente documento, EL ARRENDADOR da en arrendamiento a favor de EL ARRENDATARIO el inmueble descrito en la cláusula primera, para ser destinado única y exclusivamente para fines de {{ tipo_uso|upper }}, quedando estrictamente prohibido darle un uso distinto o subarrendarlo.')

        doc.add_heading('TERCERA: PLAZO', 2)
        p_plazo = _add_justified('El plazo de duración del contrato será de ')
        p_plazo.add_run('{{ plazo_meses }} MESES').bold = True
        p_plazo.add_run(' forzosos para ambas partes, el cual iniciará su vigencia el día ')
        p_plazo.add_run('{{ fecha_inicio }}').bold = True
        p_plazo.add_run('. Al vencimiento, el contrato podrá ser renovado previo acuerdo escrito con una anticipación mínima de treinta (30) días. Si no hay renovación, EL ARRENDATARIO deberá desocupar el inmueble inmediatamente.')

        doc.add_heading('CUARTA: RENTA Y FORMA DE PAGO', 2)
        p3 = _add_justified()
        p3.add_run('La renta mensual pactada de mutuo acuerdo asciende a la suma de S/ ')
        p3.add_run('{{ monto_numero }}').bold = True
        p3.add_run(' (')
        p3.add_run('{{ monto_texto }}').bold = True
        p3.add_run('), pagaderos por mes adelantado entre los días 1 y 5 de cada mes, mediante depósito bancario. El incumplimiento en el pago generará el interés moratorio máximo permitido por el Banco Central de Reserva del Perú.')

        doc.add_heading('QUINTA: OBLIGACIONES DEL ARRENDATARIO', 2)
        _add_justified('1. Pagar puntualmente la renta mensual acordada.\n2. Asumir el pago puntual de los servicios básicos (agua, energía eléctrica, arbitrios municipales y mantenimiento si lo hubiere).\n3. Conservar el inmueble en idénticas condiciones a las que fue recibido, no pudiendo realizar modificaciones estructurales sin la autorización expresa y por escrito de EL ARRENDADOR.\n4. Permitir la inspección del inmueble por parte de EL ARRENDADOR previa coordinación.')

        doc.add_heading('SEXTA: GARANTÍA', 2)
        p_garantia = _add_justified('A la suscripción del presente contrato, EL ARRENDATARIO entrega a favor de EL ARRENDADOR, la suma equivalente a ')
        p_garantia.add_run('{{ garantia_meses }}').bold = True
        p_garantia.add_run(' mes(es) de renta en calidad de depósito de garantía, es decir la suma de S/ ')
        p_garantia.add_run('{{ garantia_total_numero }}').bold = True
        p_garantia.add_run(' (')
        p_garantia.add_run('{{ garantia_total_texto }}').bold = True
        p_garantia.add_run('). Este monto no devengará intereses y servirá para cubrir cualquier daño al inmueble o deudas pendientes, siendo devuelto al finalizar el contrato en caso no existan obligaciones pendientes.')

        # --- CLÁUSULAS CONDICIONALES DINÁMICAS (JINJA2) ---
        _add_justified('{%p if tipo_uso == "Vivienda" %}')
        doc.add_heading('SÉPTIMA: CLÁUSULA DE ALLANAMIENTO FUTURO Y DESALOJO NOTARIAL', 2)
        _add_justified('De conformidad con el artículo 594° del Código Procesal Civil, modificado por la Ley N° 30201, EL ARRENDATARIO se allana expresamente a la demanda judicial para desocupar el inmueble por vencimiento de plazo o falta de pago. Asimismo, las partes declaran acogerse a la Ley N° 30933 que regula el Procedimiento Especial de Desalojo con Intervención Notarial.')
        _add_justified('{%p elif tipo_uso == "Comercial" %}')
        doc.add_heading('SÉPTIMA: PENALIDAD AGRESIVA POR INCUMPLIMIENTO COMERCIAL', 2)
        _add_justified('En caso de incumplimiento del pago de la renta por más de cinco (5) días calendario, EL ARRENDATARIO pagará una penalidad automática equivalente al 5% del monto de la renta por cada día de retraso. El contrato quedará resuelto de pleno derecho facultando al ARRENDADOR a tomar posesión inmediata del local.')
        _add_justified('{%p endif %}')

        doc.add_heading('NOVENA: PENALIDAD Y JURISDICCIÓN', 2)
        _add_justified('En caso EL ARRENDATARIO no desocupe el inmueble al vencimiento del plazo, pagará una penalidad diaria equivalente a S/ 100.00 (Cien y 00/100 Soles), independiente del cobro de la renta. Para cualquier controversia, las partes se someten a la competencia de los jueces y tribunales de Lima.')

        p_final = _add_justified('\nEn señal de absoluta conformidad, las partes suscriben el presente contrato por duplicado, con firmas legalizadas notarialmente en ')
        p_final.add_run('{{ fecha_actual }}').bold = True
        p_final.add_run('.')

        p_firmas = _add_justified('\n\n\n__________________________________\nEL ARRENDADOR\n')
        p_firmas.add_run('{{ nombre_arrendador }}').bold = True
        p_firmas.add_run('\nDNI: ')
        p_firmas.add_run('{{ dni_arrendador }}').bold = True
        p_firmas.add_run('\n\n\n\n__________________________________\nEL ARRENDATARIO\n')
        p_firmas.add_run('{{ nombre_arrendatario }}').bold = True
        p_firmas.add_run('\nDNI: ')
        p_firmas.add_run('{{ dni_arrendatario }}').bold = True
        p_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(self.template_path)
        logger.info("Plantilla de Arrendamiento auto-generada limpia.")

    def _crear_plantilla_demanda(self):
        """Genera el modelo estándar peruano para una demanda de alimentos."""
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Cm(2.54)
            section.left_margin = section.right_margin = Cm(2.54)
            
        def _add_justified(text=''):
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return p

        # Encabezado Judicial
        p_header = doc.add_paragraph()
        p_header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_header.add_run('Especialista: ............................\n').bold = True
        p_header.add_run('Expediente: ............................\n').bold = True
        p_header.add_run('Cuaderno: PRINCIPAL\n').bold = True
        p_header.add_run('Escrito: 01-2026\n').bold = True
        p_header.add_run('Sumilla: INTERPONGO DEMANDA DE ALIMENTOS\n').bold = True

        doc.add_heading('SEÑOR JUEZ DE PAZ LETRADO', 2)

        p = _add_justified()
        p.add_run('{{ dem_n_ddante }}').bold = True
        p.add_run(', identificada con DNI N° ')
        p.add_run('{{ dem_dni_ddante }}').bold = True
        p.add_run(', con domicilio real en ')
        p.add_run('{{ dem_dom_ddante }}').bold = True
        p.add_run(', y señalando domicilio procesal en la casilla judicial electrónica correspondiente; a usted, respetuosamente, digo:\n')

        doc.add_heading('I. DEL DEMANDADO Y SU DIRECCIÓN DOMICILIARIA:', 3)
        _add_justified('La presente demanda la dirijo contra don(ña) {{ dem_n_ddado }}, a quien se le deberá notificar con las formalidades de Ley en su domicilio ubicado en {{ dem_dom_ddado }}.')

        doc.add_heading('II. PETITORIO:', 3)
        p_petitorio = _add_justified('Que, recurro a su Despacho con la finalidad de interponer demanda de prestación de alimentos contra el demandado, solicitando se le ordene el pago de una pensión alimenticia mensual y adelantada ascendente a la suma de S/ ')
        p_petitorio.add_run('{{ dem_monto_numero }}').bold = True
        p_petitorio.add_run(' (')
        p_petitorio.add_run('{{ dem_monto_texto }}').bold = True
        p_petitorio.add_run('), a favor de nuestro menor hijo(a) alimentista de nombre ')
        p_petitorio.add_run('{{ dem_n_menor }}').bold = True
        p_petitorio.add_run(', de ')
        p_petitorio.add_run('{{ dem_edad_menor }}').bold = True
        p_petitorio.add_run(' de edad. Pensión que cubrirá los conceptos de: {{ dem_conceptos }}.')

        doc.add_heading('III. FUNDAMENTOS DE HECHO:', 3)
        _add_justified('PRIMERO.- Que, producto de las relaciones que mantuve con el demandado, procreamos a nuestro menor hijo(a) {{ dem_n_menor }}, quien actualmente se encuentra bajo mi exclusivo cuidado.')
        _add_justified('SEGUNDO.- Que, los gastos que genera la manutención son elevados. Los hechos que justifican la presente pretensión son los siguientes: {{ dem_justificacion }}.')
        _add_justified('TERCERO.- Que, el demandado se encuentra en plena capacidad física y mental para proveer económicamente a las necesidades de su hijo(a), sin embargo, viene eludiendo sus responsabilidades como padre, motivo por el cual me veo en la imperiosa necesidad de recurrir al órgano jurisdiccional.')

        doc.add_heading('IV. FUNDAMENTACIÓN JURÍDICA:', 3)
        _add_justified('Amparo mi pretensión en lo dispuesto por los artículos 472°, 473° y 481° del Código Civil Peruano, así como en el Código de los Niños y Adolescentes.')

        doc.add_heading('V. MEDIOS PROBATORIOS:', 3)
        _add_justified('1. Copia certificada del Acta de Nacimiento del menor.\n2. Documentos que acreditan los gastos (Boletas, recibos, etc.).\n3. Ficha RENIEC de las partes.')

        doc.add_heading('VI. ANEXOS:', 3)
        _add_justified('1-A. Copia de mi DNI.\n1-B. Partida de nacimiento original del menor alimentista.\n1-C. Documentos probatorios detallados en el punto V.')

        doc.add_heading('POR TANTO:', 3)
        _add_justified('A Ud., Señor Juez, solicito admitir a trámite la presente demanda y declararla fundada en su oportunidad.\n')

        p_final = _add_justified('Lima, {{ fecha_actual }}')
        p_final.alignment = WD_ALIGN_PARAGRAPH.RIGHT

        p_firmas = _add_justified('\n\n\n__________________________________\nEL DEMANDANTE\n')
        p_firmas.add_run('{{ dem_n_ddante }}').bold = True
        p_firmas.add_run('\nDNI: ')
        p_firmas.add_run('{{ dem_dni_ddante }}').bold = True
        p_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.save(self.template_path)
        logger.info(f"Plantilla de Demanda auto-generada limpia en {self.template_path}")

    def _crear_plantilla_compraventa(self):
        """Genera un contrato de compraventa con cláusulas de saneamiento y evicción (Blindado)."""
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Cm(2.54)
            section.left_margin = section.right_margin = Cm(2.54)
            
        def _add_justified(text=''):
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return p
            
        doc.add_heading('CONTRATO DE COMPRAVENTA DE BIEN INMUEBLE', 1)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        p = _add_justified('Conste por el presente documento, el Contrato de Compraventa de Bien Inmueble que celebran, de una parte: ')
        p.add_run('{{ cv_n_vend }}').bold = True
        p.add_run(', con DNI N° {{ cv_dni_vend }}, de estado civil {{ cv_est_vend }}, a quien en adelante se le denominará EL VENDEDOR; y de la otra parte, ')
        p.add_run('{{ cv_n_comp }}').bold = True
        p.add_run(', con DNI N° {{ cv_dni_comp }}, de estado civil {{ cv_est_comp }}, a quien en adelante se le denominará EL COMPRADOR, bajo las siguientes cláusulas:')
        
        doc.add_heading('PRIMERA: EL INMUEBLE', 2)
        _add_justified('EL VENDEDOR es propietario exclusivo y legítimo del inmueble ubicado en {{ cv_ubicacion }}, cuyo derecho de propiedad se encuentra debidamente inscrito en la Partida Electrónica N° {{ cv_partida }} del Registro de Predios de la Zona Registral de {{ cv_oficina }}.')
        
        doc.add_heading('SEGUNDA: TRANSFERENCIA Y PRECIO', 2)
        _add_justified('Por el presente acto, EL VENDEDOR da en venta real y enajenación perpetua el inmueble descrito en la cláusula anterior a favor de EL COMPRADOR, comprendiendo sus usos, costumbres, servidumbres, entradas, salidas y todo cuanto de hecho y por derecho le corresponda, sin reserva ni limitación alguna (Ad Corpus).')
        _add_justified('El precio pactado de mutuo acuerdo es la suma de S/ {{ cv_precio_numero }} ({{ cv_precio_texto }}), cantidad que es cancelada a la suscripción del presente documento mediante el siguiente medio de pago: {{ cv_pago }}. EL VENDEDOR declara que la firma del presente instrumento sirve como eficaz y suficiente recibo de cancelación.')
        
        doc.add_heading('TERCERA: CARGAS Y GRAVÁMENES', 2)
        _add_justified('EL VENDEDOR declara bajo juramento que sobre el inmueble materia de compraventa no pesa carga, embargo, medida cautelar, hipoteca, ni medida judicial o extrajudicial que limite su libre disposición. En todo caso, EL VENDEDOR se obliga al saneamiento por evicción y vicios ocultos, conforme a ley.')
        
        doc.add_heading('CUARTA: GASTOS Y TRIBUTOS', 2)
        _add_justified('Las partes acuerdan que todos los gastos notariales y registrales que origine la elevación a Escritura Pública del presente documento serán asumidos por EL COMPRADOR. Asimismo, EL VENDEDOR asume el pago del Impuesto Predial y Arbitrios Municipales generados hasta la fecha de suscripción de este contrato, correspondiendo a EL COMPRADOR los generados a partir de esta fecha, así como el Impuesto de Alcabala si correspondiese.')
        
        doc.add_heading('QUINTA: JURISDICCIÓN', 2)
        _add_justified('Para los efectos del presente contrato, las partes renuncian al fuero de sus domicilios y se someten expresamente a la competencia jurisdiccional de los jueces y tribunales del distrito correspondiente a la ubicación del inmueble.')
        
        _add_justified('\nSuscrito en la ciudad, en fecha {{ fecha_actual }}.')
        p_firmas = _add_justified('\n\n\n_______________________\nEL VENDEDOR\nDNI: {{ cv_dni_vend }}\n\n\n\n_______________________\nEL COMPRADOR\nDNI: {{ cv_dni_comp }}')
        p_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(self.template_path)

    def _crear_plantilla_liquidacion(self):
        """Genera una Liquidación formal y detallada (RRHH)."""
        doc = Document()
        
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        for section in doc.sections:
            section.top_margin = section.bottom_margin = Cm(2.54)
            section.left_margin = section.right_margin = Cm(2.54)
            
        def _add_justified(text=''):
            p = doc.add_paragraph(text)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            return p
            
        doc.add_heading('LIQUIDACIÓN DE BENEFICIOS SOCIALES', 1)
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        _add_justified('En cumplimiento del Texto Único Ordenado del D. Leg. 728 y la Ley de CTS, se emite el presente documento detallando el pago de beneficios laborales a favor del trabajador cesado.')
        
        doc.add_heading('I. DATOS GENERALES', 2)
        _add_justified('EMPLEADOR: {{ liq_emp }} (RUC: {{ liq_ruc }})\nTRABAJADOR: {{ liq_n_trab }} (DNI: {{ liq_dni_trab }})\nCARGO OCUPADO: {{ liq_cargo }}')
        
        doc.add_heading('II. DATOS DEL CESE', 2)
        _add_justified('FECHA DE INGRESO: {{ liq_inicio }}\nFECHA DE CESE: {{ liq_fin }}\nMOTIVO DEL CESE: {{ liq_motivo }}\nREMUNERACIÓN COMPUTABLE: S/ {{ liq_sueldo }}')
        
        doc.add_heading('III. CONCEPTOS LIQUIDADOS', 2)
        _add_justified('Conforme a ley, los conceptos liquidados en el presente acto corresponden a:\n- Compensación por Tiempo de Servicios (CTS) Trunca.\n- Gratificaciones Legales Truncas.\n- Vacaciones Truncas y Pendientes de Goce correspondientes a {{ liq_vacaciones }} días.')
        
        doc.add_heading('IV. DECLARACIÓN DE CONFORMIDAD', 2)
        _add_justified('Con la suscripción del presente documento, el trabajador declara recibir la totalidad de sus beneficios sociales a su entera satisfacción. Asimismo, deja expresa constancia de que la empresa no le adeuda suma alguna por concepto de salarios, horas extras, utilidades, indemnizaciones o cualquier otro beneficio laboral, no teniendo nada que reclamar en el futuro por vía administrativa, policial o judicial.')
        
        _add_justified('\nSuscrito en fecha {{ fecha_actual }}.')
        p_firmas = _add_justified('\n\n\n_______________________\nEL EMPLEADOR\n\n\n\n_______________________\nEL TRABAJADOR\nDNI: {{ liq_dni_trab }}')
        p_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(self.template_path)

    def _crear_plantilla_certificado(self):
        doc = Document()
        doc.add_heading('CERTIFICADO DE TRABAJO', 1)
        doc.add_paragraph('A QUIEN CORRESPONDA:')
        doc.add_paragraph('Por el presente documento, la empresa {{ cer_emp }} (RUC: {{ cer_ruc }}), certifica que el Sr./Sra. {{ cer_n_trab }}, identificado(a) con DNI N° {{ cer_dni_trab }}, ha laborado en nuestra institución desempeñando el cargo de {{ cer_cargo }}.')
        doc.add_paragraph('Durante su permanencia, cumplió principalmente con las siguientes funciones:\n{{ cer_funciones }}')
        doc.add_paragraph('\nPeriodo laborado:\nDesde: {{ cer_inicio }}\nHasta: {{ cer_fin }}')
        doc.add_paragraph('\nSe expide el presente certificado a solicitud del interesado para los fines que estime convenientes.\n\nFecha: {{ fecha_actual }}')
        p_firmas = doc.add_paragraph('\n\n\n_______________________\nGERENCIA DE RR.HH.\n{{ cer_emp }}')
        p_firmas.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.save(self.template_path)

    def _generar_sync(self, contexto: dict, output_path: str | Path) -> tuple[bool, str]:
        """
        [Interno] Lógica de renderizado sincrónico.
        """
        try:
            self._crear_plantilla_automatica()
                
            doc = DocxTemplate(self.template_path)
            doc.render(contexto)
            
            Path(output_path).parent.mkdir(parents=True, exist_ok=True)
            doc.save(output_path)
            logger.info(f"Documento generado con éxito en: {output_path}")
            return True, "Éxito"
            
        except Exception as e:
            logger.error(f"Fallo al procesar el documento Word: {e}")
            return False, str(e)

    async def generar_async(self, contexto: dict, output_path: str | Path) -> tuple[bool, str]:
        """Wrapper asíncrono para ejecutar la generación de Word sin bloquear el Event Loop."""
        return await asyncio.to_thread(self._generar_sync, contexto, output_path)